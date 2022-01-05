import docx

doc = docx.Document()
doc.add_paragraph('This is a file with a picture.')
doc.paragraphs[0].runs[0].add_break()

doc.add_picture('Word/zophie.png', width=docx.shared.Inches(1),
                height=docx.shared.Cm(4))

doc.save('Word/Added_pics.docx')
