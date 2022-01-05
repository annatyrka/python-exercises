import docx

doc = docx.Document('Word/demo.docx')

print(doc.paragraphs[0].text)

print(doc.paragraphs[0].style)
doc.paragraphs[0].style = 'Normal'

print(doc.paragraphs[1].text)
print((doc.paragraphs[1].runs[0].text, doc.paragraphs[1].runs[1].text,
       doc.paragraphs[1].runs[2].text, doc.paragraphs[1].runs[3].text,))

doc.paragraphs[1].runs[0].style = 'Quote Char'
doc.paragraphs[1].runs[1].underline = True
doc.paragraphs[1].runs[3].shadow = True

doc.save('Word/restyled_doc.docx')
