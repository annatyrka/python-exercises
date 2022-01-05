import docx

quest_list = open('Word/guests.txt')

doc = docx.Document('Word/Invitations.docx')
i = 5
for quest_name in quest_list.readlines():
    doc.add_paragraph(
        'It would be a pleasure to have the company of', 'Style1')
    doc.add_paragraph(quest_name, 'Style2')
    doc.add_paragraph('at 10010 Memory Lane on the Evening of', 'Style1')
    doc.add_paragraph('April 1st', 'Style3')
    doc.add_paragraph('at 7 o\'clock')
    doc.paragraphs[i].runs[0].add_break(docx.enum.text.WD_BREAK.PAGE)
    i += 5
doc.save('Word/Invitations.docx')
quest_list.close()
