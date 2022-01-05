import docx

doc = docx.Document()
doc.add_heading('This is my document', 1)
doc.add_paragraph('This is the first paragraph on the first page.')
doc.paragraphs[1].runs[0].add_break()
doc.add_paragraph('This is the second paragraph on the first page.')
doc.paragraphs[2].runs[0].add_break(docx.enum.text.WD_BREAK.PAGE)
doc.add_paragraph('This is a paragraph on the second page.')

doc.save('Word/Adding_breaks.docx')
