import docx

doc = docx.Document()
doc.add_paragraph('Hello, world!')
para_obj1 = doc.add_paragraph('This is a soecond paragraph.')
para_obj2 = doc.add_paragraph('This is a yet another paragraph.')
para_obj1.add_run(' This text is being added to the second paragraph.')

doc.save('Word/hello_world.docx')
